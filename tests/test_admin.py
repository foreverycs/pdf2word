"""Tests for admin console auth and pages."""

from __future__ import annotations

import importlib
from pathlib import Path

import pytest
from fastapi.testclient import TestClient


@pytest.fixture()
def admin_client(tmp_path, monkeypatch):
    d = tmp_path / "file"
    d.mkdir()
    monkeypatch.setenv("UPLOAD_FILE_DIR", str(d))
    monkeypatch.setenv("UPLOAD_RETENTION_DAYS", "5")
    monkeypatch.setenv("ALLOW_INSECURE_ADMIN", "1")
    monkeypatch.setenv("ADMIN_PASSWORD", "test-pass")
    monkeypatch.setenv("ADMIN_SECRET", "test-secret-for-unit-tests-only")

    import core.settings as settings_mod
    import core.concurrency as concurrency_mod
    import storage.history as h
    import admin.auth as auth
    import admin.routes as routes
    import admin.rate_limit as rate_limit
    import tools.common as common
    import app as app_mod

    settings_mod.clear_settings_cache()
    concurrency_mod.reset_semaphore()
    rate_limit.reset_all()
    common.refresh_limits()
    importlib.reload(auth)
    importlib.reload(routes)
    importlib.reload(app_mod)

    client = TestClient(app_mod.app)
    yield client, h, tmp_path
    rate_limit.reset_all()
    settings_mod.clear_settings_cache()


@pytest.fixture()
def hist_only(tmp_path, monkeypatch):
    d = tmp_path / "file"
    d.mkdir()
    monkeypatch.setenv("UPLOAD_FILE_DIR", str(d))
    monkeypatch.setenv("UPLOAD_RETENTION_DAYS", "5")
    monkeypatch.setenv("ALLOW_INSECURE_ADMIN", "1")
    import core.settings as settings_mod
    import storage.history as h

    settings_mod.clear_settings_cache()
    return h, tmp_path


def _csrf_token(client: TestClient) -> str:
    """Load login page so the double-submit CSRF cookie is set."""
    page = client.get("/admin/login")
    assert page.status_code == 200
    token = client.cookies.get("toolkit_csrf")
    assert token, "CSRF cookie missing after GET /admin/login"
    return token


def _login(client: TestClient, password: str = "test-pass"):
    token = _csrf_token(client)
    return client.post(
        "/admin/login",
        data={
            "password": password,
            "next": "/admin",
            "csrf_token": token,
        },
        follow_redirects=False,
    )


def test_admin_requires_login(admin_client):
    client, _, _ = admin_client
    r = client.get("/admin", follow_redirects=False)
    assert r.status_code in (303, 307, 302)
    assert "/admin/login" in r.headers.get("location", "")


def test_admin_login_and_dashboard(admin_client):
    client, _, _ = admin_client
    bad = _login(client, "wrong")
    assert bad.status_code in (303, 307, 302)
    assert "error" in bad.headers.get("location", "")

    ok = _login(client)
    assert ok.status_code in (303, 307, 302)
    loc = ok.headers.get("location", "")
    assert "/admin" in loc

    dash = client.get("/admin")
    assert dash.status_code == 200
    assert "仪表盘" in dash.text
    assert "上传" in dash.text
    assert "v0.9.0" in dash.text


def test_admin_login_requires_csrf(admin_client):
    client, _, _ = admin_client
    # No prior GET → no CSRF cookie / form token
    r = client.post(
        "/admin/login",
        data={"password": "test-pass", "next": "/admin"},
        follow_redirects=False,
    )
    assert r.status_code in (303, 307, 302)
    assert "token" in (r.headers.get("location") or "").lower()


def test_admin_login_rate_limit(admin_client, monkeypatch):
    """After repeated failures the same client is locked out briefly."""
    from urllib.parse import unquote

    import admin.rate_limit as rate_limit

    client, _, _ = admin_client
    rate_limit.reset_all()
    monkeypatch.setattr(rate_limit, "DEFAULT_MAX_FAILURES", 3)
    monkeypatch.setattr(rate_limit, "DEFAULT_LOCKOUT_SEC", 60.0)

    try:
        for _ in range(3):
            r = _login(client, "wrong-password")
            assert r.status_code in (303, 307, 302)

        locked = _login(client, "test-pass")
        assert locked.status_code in (303, 307, 302)
        loc = unquote(locked.headers.get("location", ""))
        assert "too many" in loc

        rate_limit.reset_all()
        ok = _login(client, "test-pass")
        assert ok.status_code in (303, 307, 302)
        assert "error" not in unquote(ok.headers.get("location", ""))
    finally:
        rate_limit.reset_all()


def test_admin_uploads_delete_and_download(admin_client):
    client, h, tmp_path = admin_client
    src = tmp_path / "a.pdf"
    src.write_bytes(b"%PDF-test")
    rec = h.archive_conversion(
        tool="pdf2word",
        original_name="a.pdf",
        input_path=str(src),
    )
    assert rec is not None

    _login(client)
    page = client.get("/admin/uploads")
    assert page.status_code == 200
    assert "a.pdf" in page.text
    assert rec["id"] in page.text
    assert "batch-delete" in page.text
    assert "批量删除" in page.text

    dl = client.get(f"/admin/uploads/{rec['id']}/download")
    assert dl.status_code == 200
    assert dl.content.startswith(b"%PDF")

    # CSRF cookie already set by login page + dashboard navigation
    csrf = client.cookies.get("toolkit_csrf")
    assert csrf
    deleted = client.post(
        f"/admin/uploads/{rec['id']}/delete",
        data={"csrf_token": csrf},
        follow_redirects=False,
    )
    assert deleted.status_code in (303, 307, 302)
    assert h.get_record(rec["id"]) is None


def test_admin_uploads_batch_delete(admin_client):
    client, h, tmp_path = admin_client
    recs = []
    for name in ("batch1.pdf", "batch2.pdf", "keep.pdf"):
        src = tmp_path / name
        src.write_bytes(b"%PDF-" + name.encode())
        rec = h.archive_conversion(
            tool="pdf2word",
            original_name=name,
            input_path=str(src),
        )
        assert rec is not None
        recs.append(rec)

    _login(client)
    csrf = client.cookies.get("toolkit_csrf")
    assert csrf

    from urllib.parse import unquote

    # Empty selection
    empty = client.post(
        "/admin/uploads/batch-delete",
        data={"csrf_token": csrf},
        follow_redirects=False,
    )
    assert empty.status_code in (303, 307, 302)
    assert "no selection" in unquote(empty.headers.get("location", ""))

    deleted = client.post(
        "/admin/uploads/batch-delete",
        data={
            "csrf_token": csrf,
            "ids": [recs[0]["id"], recs[1]["id"], "missing-id"],
        },
        follow_redirects=False,
    )
    assert deleted.status_code in (303, 307, 302)
    loc = deleted.headers.get("location", "")
    assert "deleted" in loc
    assert h.get_record(recs[0]["id"]) is None
    assert h.get_record(recs[1]["id"]) is None
    assert h.get_record(recs[2]["id"]) is not None


def test_storage_delete_records_batch(hist_only):
    h, tmp_path = hist_only
    ids = []
    for name in ("x.pdf", "y.pdf"):
        src = tmp_path / name
        src.write_bytes(b"data")
        rec = h.archive_conversion(
            tool="pdf2word",
            original_name=name,
            input_path=str(src),
        )
        assert rec is not None
        ids.append(rec["id"])

    assert h.delete_records([]) == 0
    assert h.delete_records(["", "  "]) == 0
    # Duplicates count once
    assert h.delete_records([ids[0], ids[0], ids[1], "nope"]) == 2
    assert h.get_record(ids[0]) is None
    assert h.get_record(ids[1]) is None


def test_admin_uploads_preview(admin_client):
    client, h, tmp_path = admin_client
    src = tmp_path / "sample.pdf"
    src.write_bytes(b"%PDF-1.4 test content")
    rec = h.archive_conversion(
        tool="pdf2word",
        original_name="sample.pdf",
        input_path=str(src),
    )
    assert rec is not None

    _login(client)

    pv = client.get(f"/admin/uploads/{rec['id']}/preview")
    assert pv.status_code == 200
    assert pv.headers.get("content-type") == "application/pdf"
    assert "inline" in pv.headers.get("content-disposition", "")
    assert pv.content.startswith(b"%PDF")

    assert client.get("/admin/uploads/nonexistent/preview").status_code == 404


def test_admin_express_list_download_delete(admin_client):
    """File express packages are managed separately under /admin/express."""
    client, _, tmp_path = admin_client
    import storage.express as ex

    src = tmp_path / "parcel.bin"
    src.write_bytes(b"express-admin-payload")
    pkg = ex.create_package(
        src,
        "parcel.bin",
        content_type="application/octet-stream",
        max_downloads=3,
        note="admin-test",
    )

    # Unauthenticated → login redirect
    denied = client.get("/admin/express", follow_redirects=False)
    assert denied.status_code in (303, 307, 302)
    assert "/admin/login" in denied.headers.get("location", "")

    _login(client)
    page = client.get("/admin/express")
    assert page.status_code == 200
    assert "文件快递" in page.text
    assert pkg["code"] in page.text
    assert "parcel.bin" in page.text
    assert "admin-test" in page.text
    assert "batch-delete" in page.text

    # Search by code
    filtered = client.get("/admin/express", params={"q": pkg["code"]})
    assert filtered.status_code == 200
    assert pkg["code"] in filtered.text

    dl = client.get(f"/admin/express/{pkg['id']}/download")
    assert dl.status_code == 200
    assert dl.content == b"express-admin-payload"

    csrf = client.cookies.get("toolkit_csrf")
    assert csrf
    deleted = client.post(
        f"/admin/express/{pkg['id']}/delete",
        data={"csrf_token": csrf},
        follow_redirects=False,
    )
    assert deleted.status_code in (303, 307, 302)
    assert "deleted" in deleted.headers.get("location", "")
    assert ex.get_package_by_id(pkg["id"]) is None


def test_admin_express_batch_delete(admin_client):
    client, _, tmp_path = admin_client
    import storage.express as ex

    pkgs = []
    for name in ("e1.txt", "e2.txt", "keep.txt"):
        src = tmp_path / name
        src.write_bytes(name.encode())
        pkgs.append(ex.create_package(src, name))

    _login(client)
    csrf = client.cookies.get("toolkit_csrf")
    assert csrf

    from urllib.parse import unquote

    empty = client.post(
        "/admin/express/batch-delete",
        data={"csrf_token": csrf},
        follow_redirects=False,
    )
    assert empty.status_code in (303, 307, 302)
    assert "no selection" in unquote(empty.headers.get("location", ""))

    deleted = client.post(
        "/admin/express/batch-delete",
        data={
            "csrf_token": csrf,
            "ids": [pkgs[0]["id"], pkgs[1]["id"], "missing-id"],
        },
        follow_redirects=False,
    )
    assert deleted.status_code in (303, 307, 302)
    assert "deleted" in deleted.headers.get("location", "")
    assert ex.get_package_by_id(pkgs[0]["id"]) is None
    assert ex.get_package_by_id(pkgs[1]["id"]) is None
    assert ex.get_package_by_id(pkgs[2]["id"]) is not None


def test_admin_express_list_download_delete(admin_client):
    """File express packages are managed separately under /admin/express."""
    client, _, tmp_path = admin_client
    import storage.express as ex

    src = tmp_path / "parcel.bin"
    src.write_bytes(b"express-admin-payload")
    pkg = ex.create_package(
        src,
        "parcel.bin",
        content_type="application/octet-stream",
        max_downloads=3,
        note="admin-test",
    )

    # Unauthenticated → login redirect
    denied = client.get("/admin/express", follow_redirects=False)
    assert denied.status_code in (303, 307, 302)
    assert "/admin/login" in denied.headers.get("location", "")

    _login(client)
    page = client.get("/admin/express")
    assert page.status_code == 200
    assert "文件快递" in page.text
    assert pkg["code"] in page.text
    assert "parcel.bin" in page.text
    assert "admin-test" in page.text
    assert "batch-delete" in page.text

    # Search by code
    filtered = client.get("/admin/express", params={"q": pkg["code"]})
    assert filtered.status_code == 200
    assert pkg["code"] in filtered.text

    dl = client.get(f"/admin/express/{pkg['id']}/download")
    assert dl.status_code == 200
    assert dl.content == b"express-admin-payload"

    csrf = client.cookies.get("toolkit_csrf")
    assert csrf
    deleted = client.post(
        f"/admin/express/{pkg['id']}/delete",
        data={"csrf_token": csrf},
        follow_redirects=False,
    )
    assert deleted.status_code in (303, 307, 302)
    assert "deleted" in deleted.headers.get("location", "")
    assert ex.get_package_by_id(pkg["id"]) is None


def test_admin_express_batch_delete(admin_client):
    client, _, tmp_path = admin_client
    import storage.express as ex

    pkgs = []
    for name in ("e1.txt", "e2.txt", "keep.txt"):
        src = tmp_path / name
        src.write_bytes(name.encode())
        pkgs.append(ex.create_package(src, name))

    _login(client)
    csrf = client.cookies.get("toolkit_csrf")
    assert csrf

    from urllib.parse import unquote

    empty = client.post(
        "/admin/express/batch-delete",
        data={"csrf_token": csrf},
        follow_redirects=False,
    )
    assert empty.status_code in (303, 307, 302)
    assert "no selection" in unquote(empty.headers.get("location", ""))

    deleted = client.post(
        "/admin/express/batch-delete",
        data={
            "csrf_token": csrf,
            "ids": [pkgs[0]["id"], pkgs[1]["id"], "missing-id"],
        },
        follow_redirects=False,
    )
    assert deleted.status_code in (303, 307, 302)
    assert "deleted" in deleted.headers.get("location", "")
    assert ex.get_package_by_id(pkgs[0]["id"]) is None
    assert ex.get_package_by_id(pkgs[1]["id"]) is None
    assert ex.get_package_by_id(pkgs[2]["id"]) is not None


def test_admin_uploads_word_preview_pdf(admin_client, monkeypatch):
    """Word .docx/.doc preview converts to PDF via word2pdf (mocked)."""
    from docx import Document

    client, h, tmp_path = admin_client
    src = tmp_path / "note.docx"
    doc = Document()
    doc.add_paragraph("Hello Word preview")
    doc.save(str(src))

    rec = h.archive_conversion(
        tool="word2pdf",
        original_name="note.docx",
        input_path=str(src),
    )
    assert rec is not None

    def fake_convert(input_path, output_path=None, **kwargs):
        out = Path(output_path or str(Path(input_path).with_suffix(".pdf")))
        out.write_bytes(b"%PDF-1.4 fake preview")
        return str(out), "libreoffice"

    monkeypatch.setattr("word2pdf.convert_to_pdf", fake_convert)

    _login(client)
    pv = client.get(f"/admin/uploads/{rec['id']}/preview")
    assert pv.status_code == 200
    assert "application/pdf" in (pv.headers.get("content-type") or "")
    assert pv.headers.get("x-preview-source") == "word"
    assert "inline" in (pv.headers.get("content-disposition") or "")
    assert pv.content.startswith(b"%PDF")

    # Second request should hit disk cache (no convert call needed).
    calls = {"n": 0}

    def boom(*a, **k):
        calls["n"] += 1
        raise AssertionError("should use cache")

    monkeypatch.setattr("word2pdf.convert_to_pdf", boom)
    pv2 = client.get(f"/admin/uploads/{rec['id']}/preview")
    assert pv2.status_code == 200
    assert calls["n"] == 0
    assert pv2.content.startswith(b"%PDF")


def test_admin_uploads_word_preview_chinese_filename(admin_client, monkeypatch):
    """Non-ASCII original names must not 500 on Content-Disposition."""
    from docx import Document

    client, h, tmp_path = admin_client
    src = tmp_path / "cn.docx"
    doc = Document()
    doc.add_paragraph("中文内容预览")
    doc.save(str(src))

    rec = h.archive_conversion(
        tool="word2pdf",
        original_name="中文报告.docx",
        input_path=str(src),
    )
    assert rec is not None

    def fake_convert(input_path, output_path=None, **kwargs):
        out = Path(output_path or str(Path(input_path).with_suffix(".pdf")))
        out.write_bytes(b"%PDF-1.4 cn preview")
        return str(out), "libreoffice"

    monkeypatch.setattr("word2pdf.convert_to_pdf", fake_convert)

    _login(client)
    pv = client.get(f"/admin/uploads/{rec['id']}/preview")
    assert pv.status_code == 200, pv.text[:500]
    assert "application/pdf" in (pv.headers.get("content-type") or "")
    cd = pv.headers.get("content-disposition") or ""
    assert "inline" in cd
    cd.encode("latin-1")
    assert "filename*=" in cd


def test_admin_uploads_word_preview_no_engine(admin_client, monkeypatch):
    """Missing conversion engine returns 503."""
    from core.errors import EngineNotFoundError
    from docx import Document

    client, h, tmp_path = admin_client
    src = tmp_path / "note.docx"
    Document().save(str(src))
    rec = h.archive_conversion(
        tool="word2pdf",
        original_name="note.docx",
        input_path=str(src),
    )
    assert rec is not None

    def no_engine(*a, **k):
        raise EngineNotFoundError("No conversion engine available")

    monkeypatch.setattr("word2pdf.convert_to_pdf", no_engine)

    _login(client)
    pv = client.get(f"/admin/uploads/{rec['id']}/preview")
    assert pv.status_code == 503
    assert "engine" in (pv.json().get("detail") or "").lower()


def test_admin_api_stats_unauthorized(admin_client):
    client, _, _ = admin_client
    r = client.get("/admin/api/stats")
    assert r.status_code == 401


def test_admin_api_stats_ok(admin_client):
    client, _, _ = admin_client
    _login(client)
    r = client.get("/admin/api/stats")
    assert r.status_code == 200
    body = r.json()
    assert "storage" in body
    assert "health" in body


def test_storage_delete_and_stats(hist_only):
    h, tmp_path = hist_only
    src = tmp_path / "b.pdf"
    src.write_bytes(b"data")
    rec = h.archive_conversion(
        tool="pdf2word",
        original_name="b.pdf",
        input_path=str(src),
        extra={"pages": 1},
    )
    assert rec is not None
    stats = h.storage_stats()
    assert stats["record_count"] >= 1
    assert "pdf2word" in stats["by_tool"]
    assert h.get_record(rec["id"]) is not None
    assert h.delete_record(rec["id"]) is True
    assert h.get_record(rec["id"]) is None
    assert h.delete_record("no-such") is False
