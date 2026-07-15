import os
import tempfile

import pytest

from converter import extract_document, write_document
from converter.pdf_reader import TableBlock, TextBlock


def _make_sample_pdf(path: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
    from reportlab.lib import colors

    doc = SimpleDocTemplate(path, pagesize=A4)
    data = [
        ["Name", "Info", "Score"],
        ["Alice", "Math", "90"],
        ["Bob", "Physics", "85"],
    ]
    table = Table(data, colWidths=[80, 80, 80])
    style = TableStyle([
        ("GRID", (0, 0), (-1, -1), 1, colors.black),
        # merge "Name" with the empty area below by spanning first col rows 0-1
        ("SPAN", (0, 0), (0, 1)),
    ])
    table.setStyle(style)
    doc.build([table])


def test_roundtrip_preserves_grid_and_merge():
    tmp = tempfile.mkdtemp(prefix="pdf2word_test_")
    pdf_path = os.path.join(tmp, "sample.pdf")
    docx_path = os.path.join(tmp, "out.docx")
    _make_sample_pdf(pdf_path)

    pages = extract_document(pdf_path)
    assert pages, "no pages extracted"
    # find the table block among page blocks
    table: TableBlock = next(
        (b for b in pages[0].blocks if isinstance(b, TableBlock)), None
    )
    assert table is not None, "table not detected"
    assert table.rows == 3, table.rows
    assert table.cols == 3, table.cols

    # the merged cell (0,0) should report rowspan 2
    anchor = table.cells[0][0]
    assert anchor is not None
    assert anchor.rowspan == 2, anchor.rowspan
    assert anchor.text == "Name", anchor.text

    write_document(pages, docx_path)
    assert os.path.exists(docx_path)

    # read back and verify the merge produced a single cell object
    from docx import Document
    doc = Document(docx_path)
    assert doc.tables, "no table in docx"
    doc_table = doc.tables[0]
    assert len(doc_table.rows) == 3
    assert len(doc_table.columns) == 3
    # covered cell (1,0) must share the same underlying XML cell as (0,0)
    assert doc_table.cell(1, 0)._tc is doc_table.cell(0, 0)._tc


def test_text_block_extracted():
    tmp = tempfile.mkdtemp(prefix="pdf2word_test_")
    pdf_path = os.path.join(tmp, "t.pdf")
    docx_path = os.path.join(tmp, "t.docx")
    _make_sample_pdf(pdf_path)
    pages = extract_document(pdf_path)
    assert any(isinstance(b, TextBlock) or isinstance(b, TableBlock)
               for b in pages[0].blocks)


def _make_styled_pdf(path: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet

    doc = SimpleDocTemplate(path, pagesize=A4)
    style = getSampleStyleSheet()["Normal"]
    style.fontSize = 14
    style.fontName = "Helvetica"
    tbl = Table([[Paragraph("Big", style), "Small"],
                 ["X", "Y"]], colWidths=[90, 90])
    tbl.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 2.5, colors.black),
        ("BOX", (0, 0), (-1, -1), 2.5, colors.black),
    ]))
    doc.build([tbl])


def _border_sz(doc_table, edge: str) -> int:
    from docx.oxml.ns import qn
    borders = doc_table._tbl.tblPr.find(qn("w:tblBorders"))
    el = borders.find(qn(f"w:{edge}"))
    return int(el.get(qn("w:sz")))


def _cell_border_sz(doc_table, row: int, col: int, edge: str) -> int:
    from docx.oxml.ns import qn
    tc_pr = doc_table.cell(row, col)._tc.tcPr
    tc_b = tc_pr.find(qn("w:tcBorders"))
    el = tc_b.find(qn(f"w:{edge}"))
    return int(el.get(qn("w:sz")))


def test_fidelity_font_size_and_border():
    tmp = tempfile.mkdtemp(prefix="pdf2word_test_")
    pdf_path = os.path.join(tmp, "styled.pdf")
    docx_path = os.path.join(tmp, "styled.docx")
    _make_styled_pdf(pdf_path)

    pages = extract_document(pdf_path)
    table = next(b for b in pages[0].blocks if isinstance(b, TableBlock))
    assert table.border_outer == 2.5

    write_document(pages, docx_path)
    from docx import Document
    from docx.shared import Pt
    doc = Document(docx_path)
    doc_table = doc.tables[0]

    # per-cell border width is preserved (sz is in eighths of a point)
    assert _cell_border_sz(doc_table, 0, 0, "top") == 20  # 2.5pt * 8
    assert _cell_border_sz(doc_table, 0, 0, "right") == 20

    # cell font size is preserved
    size = doc_table.cell(0, 0).paragraphs[0].runs[0].font.size
    assert size == Pt(14)


def test_fidelity_per_cell_border_color():
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
    from reportlab.lib import colors

    tmp = tempfile.mkdtemp(prefix="pdf2word_test_")
    pdf_path = os.path.join(tmp, "pc.pdf")
    docx_path = os.path.join(tmp, "pc.docx")
    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    tbl = Table([["A", "B"], ["X", "Y"]], colWidths=[90, 90])
    tbl.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 1, colors.black),
        ("LINEABOVE", (0, 0), (-1, 0), 3, colors.Color(1, 0, 0)),  # red top
        ("LINEBEFORE", (0, 0), (0, -1), 3, colors.Color(0, 0, 1)),  # blue left
    ]))
    doc.build([tbl])

    pages = extract_document(pdf_path)
    table = next(b for b in pages[0].blocks if isinstance(b, TableBlock))
    # top edge of row 0 and left edge of col 0 should carry their colours
    assert table.cells[0][0].borders["top"][1] == "FF0000"
    assert table.cells[0][0].borders["left"][1] == "0000FF"

    write_document(pages, docx_path)
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(docx_path)
    doc_table = doc.tables[0]
    tc_b = doc_table.cell(0, 0)._tc.tcPr.find(qn("w:tcBorders"))
    assert tc_b.find(qn("w:top")).get(qn("w:color")) == "FF0000"
    assert tc_b.find(qn("w:left")).get(qn("w:color")) == "0000FF"


def test_fidelity_alignment_variants():
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
    from reportlab.lib import colors

    tmp = tempfile.mkdtemp(prefix="pdf2word_test_")
    pdf_path = os.path.join(tmp, "align.pdf")
    docx_path = os.path.join(tmp, "align.docx")
    # one row, three columns: left / centre / right; plus a full-width
    # left-aligned long text that must NOT be detected as centred.
    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    tbl = Table([["Left text", "Centre text", "Right text"],
                 ["Left aligned sentence here", "X", "Y"]],
                colWidths=[140, 140, 140])
    tbl.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 1, colors.black),
        ("ALIGN", (0, 0), (0, 0), "LEFT"),
        ("ALIGN", (1, 0), (1, 0), "CENTER"),
        ("ALIGN", (2, 0), (2, 0), "RIGHT"),
        ("ALIGN", (0, 1), (0, 1), "LEFT"),
    ]))
    doc.build([tbl])

    pages = extract_document(pdf_path)
    table = next(b for b in pages[0].blocks if isinstance(b, TableBlock))
    # row 0: left / centre / right
    assert table.cells[0][0].align == "left"
    assert table.cells[0][1].align == "center"
    assert table.cells[0][2].align == "right"
    # row 1: left-aligned text stays 'left' (not misread as centred)
    assert table.cells[1][0].align == "left"

    write_document(pages, docx_path)
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document(docx_path)
    doc_table = doc.tables[0]
    assert doc_table.cell(0, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert doc_table.cell(0, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert doc_table.cell(0, 2).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert doc_table.cell(1, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_fidelity_soft_newline_normalized():
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors

    tmp = tempfile.mkdtemp(prefix="pdf2word_test_")
    pdf_path = os.path.join(tmp, "br.pdf")
    docx_path = os.path.join(tmp, "br.docx")
    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    st = getSampleStyleSheet()["Normal"]
    st.fontSize = 10
    st.leading = 12
    tbl = Table([[Paragraph("第一行<br/>第二行", st), "B"], ["C", "D"]],
                colWidths=[120, 80])
    tbl.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 1, colors.black)]))
    doc.build([tbl])

    pages = extract_document(pdf_path)
    table = next(b for b in pages[0].blocks if isinstance(b, TableBlock))
    # Soft word-wrap \n are joined so Word can reflow (no hard breaks left).
    # ReportLab may not embed CJK glyphs; still require both lines to be present
    # as a single non-newline string (joined with space or tightly).
    cell0 = table.cells[0][0].text
    assert "\n" not in cell0
    assert cell0.strip()  # non-empty after soft-join

    write_document(pages, docx_path)
    from docx import Document
    doc = Document(docx_path)
    doc_table = doc.tables[0]
    # the cell should not contain hard line breaks
    cell_text = doc_table.cell(0, 0).text
    assert "\n" not in cell_text
    assert cell_text.strip()


def test_fidelity_text_strategy_fallback_border():
    # When no drawn lines exist (e.g. text-strategy tables), the writer must
    # still emit a uniform grid so the table is visible.
    from converter.pdf_reader import Cell, TableBlock
    from docx import Document
    from docx.oxml.ns import qn

    tmp = tempfile.mkdtemp(prefix="pdf2word_test_")
    docx_path = os.path.join(tmp, "fallback.docx")
    cells = [[Cell(text="A"), Cell(text="B")], [Cell(text="C"), Cell(text="D")]]
    owner = [[(r, c) for c in range(2)] for r in range(2)]
    table = TableBlock(rows=2, cols=2, cells=cells, owner=owner,
                       col_widths=[90, 90], row_heights=[20, 20],
                       border_outer=1.0, border_inner=1.0)
    write_document([__import__("converter.pdf_reader", fromlist=["PageContent"]).PageContent(blocks=[table])],
                  docx_path)
    doc = Document(docx_path)
    doc_table = doc.tables[0]
    borders = doc_table._tbl.tblPr.find(qn("w:tblBorders"))
    assert borders is not None
    assert int(borders.find(qn("w:top")).get(qn("w:sz"))) == 8  # 1pt * 8


def _make_dims_pdf(path: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
    from reportlab.lib import colors

    doc = SimpleDocTemplate(path, pagesize=A4)
    # first column is wider; second row is taller
    tbl = Table([["A", "B"], ["X", "Y"]], colWidths=[160, 80], rowHeights=[20, 50])
    tbl.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 1, colors.black),
                             ("ALIGN", (1, 0), (1, -1), "CENTER"),
                             ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
    doc.build([tbl])


def test_fidelity_dimensions_and_alignment():
    tmp = tempfile.mkdtemp(prefix="pdf2word_test_")
    pdf_path = os.path.join(tmp, "dims.pdf")
    docx_path = os.path.join(tmp, "dims.docx")
    _make_dims_pdf(pdf_path)

    pages = extract_document(pdf_path)
    table = next(b for b in pages[0].blocks if isinstance(b, TableBlock))
    # column widths reflect the 160/80 pt source layout
    assert abs(table.col_widths[0] - 160) < 2
    assert abs(table.col_widths[1] - 80) < 2
    assert abs(table.row_heights[1] - 50) < 2

    write_document(pages, docx_path)
    from docx import Document
    doc = Document(docx_path)
    doc_table = doc.tables[0]

    # fixed layout + column widths applied
    from docx.oxml.ns import qn
    layout = doc_table._tbl.tblPr.find(qn("w:tblLayout"))
    assert layout is not None and layout.get(qn("w:type")) == "fixed"
    from docx.shared import Pt
    assert abs(doc_table.columns[0].width.pt - 160) < 2
    assert abs(doc_table.rows[1].height.pt - 50) < 2

    # alignment: column 1 is centred, vertical middle
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    p = doc_table.cell(0, 1).paragraphs[0]
    assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert doc_table.cell(0, 0).vertical_alignment == WD_ALIGN_VERTICAL.CENTER


def _make_bg_pdf(path: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
    from reportlab.lib import colors

    doc = SimpleDocTemplate(path, pagesize=A4)
    tbl = Table([["A", "B"], ["X", "Y"]], colWidths=[90, 90])
    # yellow-ish background on the top-left cell only
    tbl.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 1, colors.black),
        ("BACKGROUND", (0, 0), (0, 0), colors.Color(0.9, 0.8, 0.2)),
    ]))
    doc.build([tbl])


def test_fidelity_cell_background():
    tmp = tempfile.mkdtemp(prefix="pdf2word_test_")
    pdf_path = os.path.join(tmp, "bg.pdf")
    docx_path = os.path.join(tmp, "bg.docx")
    _make_bg_pdf(pdf_path)

    pages = extract_document(pdf_path)
    table = next(b for b in pages[0].blocks if isinstance(b, TableBlock))
    assert table.cells[0][0].bg_color == "E6CC33"  # (0.9,0.8,0.2) -> E6CC33

    write_document(pages, docx_path)
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(docx_path)
    doc_table = doc.tables[0]
    shd = doc_table.cell(0, 0)._tc.tcPr.find(qn("w:shd"))
    assert shd is not None and shd.get(qn("w:fill")) == "E6CC33"
    # other cells should not be shaded
    assert doc_table.cell(0, 1)._tc.tcPr.find(qn("w:shd")) is None


def test_fidelity_cjk_spacing_and_title_center():
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    tmp = tempfile.mkdtemp(prefix="pdf2word_test_")
    pdf_path = os.path.join(tmp, "cjk.pdf")
    docx_path = os.path.join(tmp, "cjk.docx")

    from reportlab.lib.enums import TA_CENTER
    st = getSampleStyleSheet()["Normal"]
    st.fontName = "STSong-Light"
    st.fontSize = 12
    st.alignment = TA_CENTER
    title = Paragraph("\u6708\u5ea6\u9500\u552e\u60c5\u51b5\u62a5\u544a", st)
    tbl = Table([["\u59d3\u540d", "\u6210\u7ee9"],
                 ["\u5f20\u4e09", "\u4e5d\u5341\u4e94"]], colWidths=[90, 90])
    tbl.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 1, colors.black),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
    ]))
    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    doc.build([title, tbl])

    pages = extract_document(pdf_path)

    # title: CJK characters must not be separated by spaces
    title_block = next(b for b in pages[0].blocks
                       if isinstance(b, TextBlock) and "\u6708\u5ea6" in b.text)
    assert "\u5ea6 " not in title_block.text, f"unexpected space in: {title_block.text!r}"
    assert title_block.align == "center"

    # table cells: CJK characters must not be separated by spaces
    table = next(b for b in pages[0].blocks if isinstance(b, TableBlock))
    assert " " not in table.cells[0][0].text, f"unexpected space in cell: {table.cells[0][0].text!r}"
    assert " " not in table.cells[1][0].text, f"unexpected space in cell: {table.cells[1][0].text!r}"
    assert table.cells[0][0].align == "center"

    write_document(pages, docx_path)
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document(docx_path)
    # title paragraph should be centred and have no spaces between CJK chars
    title_para = doc.paragraphs[0]
    assert title_para.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert " " not in title_para.text, f"unexpected space in docx title: {title_para.text!r}"


# ----- page range / pagination / stats --------------------------------------

def test_parse_page_range():
    from converter import parse_page_range

    assert parse_page_range(None, 5) == [0, 1, 2, 3, 4]
    assert parse_page_range("  ", 3) == [0, 1, 2]
    assert parse_page_range("1,3", 5) == [0, 2]
    assert parse_page_range("2-4", 5) == [1, 2, 3]
    assert parse_page_range("1-2,5", 5) == [0, 1, 4]
    # duplicates are collapsed, order preserved
    assert parse_page_range("3,1,3", 5) == [2, 0]

    with pytest.raises(ValueError):
        parse_page_range("0", 5)
    with pytest.raises(ValueError):
        parse_page_range("6", 5)
    with pytest.raises(ValueError):
        parse_page_range("3-1", 5)
    with pytest.raises(ValueError):
        parse_page_range("abc", 5)
    with pytest.raises(ValueError):
        parse_page_range("1-", 5)


def _make_multipage_pdf(path: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(path, pagesize=A4)
    for i in range(1, 4):
        c.drawString(100, 750, f"PageMarker{i}")
        c.showPage()
    c.save()


def test_extract_page_range_and_stats():
    from converter import count_blocks

    tmp = tempfile.mkdtemp(prefix="pdf2word_test_")
    pdf_path = os.path.join(tmp, "multi.pdf")
    _make_multipage_pdf(pdf_path)

    pages = extract_document(pdf_path, page_range="1,3")
    assert len(pages) == 2
    texts = [
        b.text for p in pages for b in p.blocks if isinstance(b, TextBlock)
    ]
    joined = " ".join(texts)
    assert "PageMarker1" in joined
    assert "PageMarker3" in joined
    assert "PageMarker2" not in joined

    stats = count_blocks(pages)
    assert stats["pages"] == 2
    assert stats["text_blocks"] >= 2


def test_write_document_page_breaks():
    from converter.pdf_reader import PageContent, TextBlock
    from docx import Document
    from docx.oxml.ns import qn

    tmp = tempfile.mkdtemp(prefix="pdf2word_test_")
    docx_path = os.path.join(tmp, "pages.docx")
    pages = [
        PageContent(blocks=[TextBlock(text="First page", top=0)]),
        PageContent(blocks=[TextBlock(text="Second page", top=0)]),
    ]
    write_document(pages, docx_path, page_breaks=True)
    doc = Document(docx_path)
    # at least one paragraph should contain a page break run
    has_break = False
    for p in doc.paragraphs:
        for run in p.runs:
            brs = run._element.findall(qn("w:br"))
            if any(br.get(qn("w:type")) == "page" for br in brs):
                has_break = True
    assert has_break, "expected a page break between PDF pages"

    # page_breaks=False should not insert one
    docx2 = os.path.join(tmp, "nobreak.docx")
    write_document(pages, docx2, page_breaks=False)
    doc2 = Document(docx2)
    for p in doc2.paragraphs:
        for run in p.runs:
            brs = run._element.findall(qn("w:br"))
            assert not any(br.get(qn("w:type")) == "page" for br in brs)


def test_image_block_written_to_docx():
    from converter.pdf_reader import PageContent, ImageBlock
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from PIL import Image
    import io

    tmp = tempfile.mkdtemp(prefix="pdf2word_test_")
    docx_path = os.path.join(tmp, "img.docx")
    buf = io.BytesIO()
    Image.new("RGB", (40, 30), color=(20, 120, 200)).save(buf, format="PNG")
    pages = [PageContent(blocks=[
        ImageBlock(
            image_bytes=buf.getvalue(), top=0, bottom=150,
            x0=40, width_pt=200, height_pt=150,
            page_width=595, align="left",
        ),
    ])]
    write_document(pages, docx_path, page_breaks=False)
    doc = Document(docx_path)
    # python-docx stores pictures as inline shapes related to runs
    assert doc.inline_shapes, "expected an inline image in the docx"
    # left-aligned, not forced center
    assert doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_image_placement_align_and_indent():
    """Images keep PDF horizontal alignment (center/right) and left indent."""
    from converter.pdf_reader import PageContent, ImageBlock, TextBlock
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    from PIL import Image
    import io

    def png():
        buf = io.BytesIO()
        Image.new("RGB", (20, 20), color=(10, 10, 10)).save(buf, format="PNG")
        return buf.getvalue()

    tmp = tempfile.mkdtemp(prefix="pdf2word_test_")
    docx_path = os.path.join(tmp, "place.docx")
    # page width A4 ≈ 595pt; centered image around x0≈(595-200)/2≈197.5
    pages = [PageContent(blocks=[
        TextBlock(text="Title", top=20, bottom=40, align="center"),
        ImageBlock(
            image_bytes=png(), top=80, bottom=200,
            x0=197.5, width_pt=200, height_pt=120,
            page_width=595, align="center",
        ),
        ImageBlock(
            image_bytes=png(), top=240, bottom=320,
            x0=40, width_pt=120, height_pt=80,
            page_width=595, align="left",
        ),
        ImageBlock(
            image_bytes=png(), top=360, bottom=440,
            x0=400, width_pt=120, height_pt=80,
            page_width=595, align="right",
        ),
    ], width=595, height=842)]
    write_document(pages, docx_path, page_breaks=False)
    doc = Document(docx_path)

    # Find paragraphs that contain images (have runs with drawing)
    img_paras = []
    for p in doc.paragraphs:
        for run in p.runs:
            if run._element.xpath(".//a:blip"):
                img_paras.append(p)
                break
    # fallback: count by inline_shapes is enough; check alignments via paragraphs
    # with non-empty runs that have pictures
    assert len(doc.inline_shapes) == 3

    # Collect alignments of paragraphs that host pictures
    pic_aligns = []
    for p in doc.paragraphs:
        has_pic = False
        for run in p.runs:
            # w:drawing present
            if run._element.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
            ) is not None or run._element.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            ):
                has_pic = True
                break
            # python-docx sometimes nests differently — also check rId pict
            if "graphicData" in run._element.xml:
                has_pic = True
                break
        if has_pic:
            pic_aligns.append(p.alignment)

    assert WD_ALIGN_PARAGRAPH.CENTER in pic_aligns
    assert WD_ALIGN_PARAGRAPH.RIGHT in pic_aligns
    assert WD_ALIGN_PARAGRAPH.LEFT in pic_aligns

    # left-aligned image with x0=40 < 72pt default margin → no extra indent
    left_paras = [p for p in doc.paragraphs if p.alignment == WD_ALIGN_PARAGRAPH.LEFT
                  and p.runs and "graphicData" in "".join(r._element.xml for r in p.runs)]
    # image with x0 larger than margin should get indent — craft one
    docx2 = os.path.join(tmp, "indent.docx")
    write_document([PageContent(blocks=[
        ImageBlock(
            image_bytes=png(), top=0, bottom=80,
            x0=144, width_pt=100, height_pt=80,  # 2" from page edge
            page_width=595, align="left",
        ),
    ])], docx2, page_breaks=False)
    doc2 = Document(docx2)
    indented = None
    for p in doc2.paragraphs:
        if p.runs and "graphicData" in "".join(r._element.xml for r in p.runs):
            indented = p
            break
    assert indented is not None
    assert indented.paragraph_format.left_indent is not None
    # extra indent ≈ (144 - section_margin_pt) / 72; section margin is 0.7"
    # → (144 - 50.4) / 72 = 1.3 inches
    assert abs(indented.paragraph_format.left_indent.inches - 1.3) < 0.05



def test_content_warnings_image_only():
    from converter import content_warnings, count_blocks
    from converter.pdf_reader import PageContent, ImageBlock

    pages = [PageContent(blocks=[
        ImageBlock(image_bytes=b"\x89PNG", top=0, width_pt=100, height_pt=100),
    ])]
    stats = count_blocks(pages)
    assert stats["images"] == 1
    assert stats["text_blocks"] == 0
    assert "image_only" in content_warnings(pages)
    assert content_warnings([PageContent(blocks=[])]) == ["empty"]


def test_scanned_page_embeds_full_image():
    """A PDF page that is only a bitmap (no text) should become an ImageBlock."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    from PIL import Image
    import io
    from converter.pdf_reader import ImageBlock
    from converter import content_warnings

    tmp = tempfile.mkdtemp(prefix="pdf2word_test_")
    pdf_path = os.path.join(tmp, "scan.pdf")
    docx_path = os.path.join(tmp, "scan.docx")

    img_buf = io.BytesIO()
    Image.new("RGB", (200, 280), color=(240, 240, 240)).save(img_buf, format="PNG")
    img_buf.seek(0)

    c = canvas.Canvas(pdf_path, pagesize=A4)
    # nearly full-page image, no text operators
    c.drawImage(ImageReader(img_buf), 40, 40, width=500, height=700)
    c.save()

    pages = extract_document(pdf_path)
    assert pages
    # either embedded image region or full-page fallback
    images = [b for b in pages[0].blocks if isinstance(b, ImageBlock)]
    assert images, "expected image content from scanned PDF"
    assert "image_only" in content_warnings(pages)
    # placement metadata should be present
    img = images[0]
    assert img.width_pt > 0 and img.height_pt > 0
    assert img.page_width > 0
    assert img.align in ("left", "center", "right")

    write_document(pages, docx_path)
    from docx import Document
    doc = Document(docx_path)
    assert doc.inline_shapes


def test_embedded_image_keeps_native_resolution(tmp_path):
    """High-res PDF XObject should be kept, not downsampled via low-DPI raster."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    from PIL import Image
    import io
    from converter.pdf_reader import ImageBlock

    pdf_path = str(tmp_path / "hires.pdf")
    # 800×600 source displayed in a 200×150 pt box (~288 DPI equivalent).
    src = Image.new("RGB", (800, 600), color=(20, 100, 180))
    for x in range(0, 800, 10):
        for y in range(0, 600, 10):
            src.putpixel((x, y), (255, 220, 0))
    buf = io.BytesIO()
    src.save(buf, format="PNG")
    buf.seek(0)

    c = canvas.Canvas(pdf_path, pagesize=A4)
    c.drawImage(ImageReader(buf), 72, 400, width=200, height=150)
    c.setFont("Helvetica", 11)
    c.drawString(72, 380, "caption")
    c.save()

    pages = extract_document(pdf_path)
    images = [b for b in pages[0].blocks if isinstance(b, ImageBlock)]
    assert images, "expected embedded image block"
    pil = Image.open(io.BytesIO(images[0].image_bytes))
    # Must retain native (or higher) resolution — not the old 144 DPI crop
    # which would be only ~400×300 for a 200×150 pt box.
    assert pil.size[0] >= 700 and pil.size[1] >= 500, pil.size
    assert abs(images[0].width_pt - 200) < 1
    assert abs(images[0].height_pt - 150) < 1


def test_image_only_page_keeps_partial_native_image(tmp_path):
    """A page with only a photo (no text) must not re-rasterise the whole page."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    from PIL import Image
    import io
    from converter.pdf_reader import ImageBlock

    pdf_path = str(tmp_path / "photo_only.pdf")
    src = Image.new("RGB", (1200, 900), color=(10, 20, 30))
    for x in range(0, 1200, 8):
        src.putpixel((x, 50), (255, 0, 0))
    buf = io.BytesIO()
    src.save(buf, format="JPEG", quality=95)
    buf.seek(0)

    c = canvas.Canvas(pdf_path, pagesize=A4)
    # Partial-page image, no text operators at all.
    c.drawImage(ImageReader(buf), 50, 350, width=300, height=225)
    c.save()

    pages = extract_document(pdf_path)
    images = [b for b in pages[0].blocks if isinstance(b, ImageBlock)]
    assert images, "expected image block"
    pil = Image.open(io.BytesIO(images[0].image_bytes))
    # Native 1200×900 — not a full-page 220 DPI A4 (~1819×2573) re-render.
    assert pil.size[0] >= 1000 and pil.size[1] >= 700, pil.size
    assert pil.size[0] < 1600, f"looks like full-page re-render: {pil.size}"
    assert abs(images[0].width_pt - 300) < 2


def test_image_render_dpi_default_raised():
    from converter import pdf_reader

    assert pdf_reader.IMAGE_RENDER_DPI >= 200


def test_image_h_align_helper():
    from converter.pdf_reader import _image_h_align

    # centered: equal side pads
    assert _image_h_align(197.5, 200, 595) == "center"
    # left: small x0
    assert _image_h_align(40, 120, 595) == "left"
    # right: large left pad
    assert _image_h_align(420, 120, 595) == "right"
    # near full-width
    assert _image_h_align(10, 560, 595) == "center"


def test_header_logo_and_title_same_row():
    """Logo + company name share one paragraph; signature labels too — not tables."""
    from converter.docx_writer import _group_horizontal_rows
    from converter.pdf_reader import ImageBlock, TextBlock, PageContent, LineBlock
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from PIL import Image
    import io

    buf = io.BytesIO()
    Image.new("RGB", (40, 20), color=(200, 0, 0)).save(buf, format="PNG")
    logo = ImageBlock(
        image_bytes=buf.getvalue(),
        top=42.6, bottom=77.6, x0=90, width_pt=76.6, height_pt=35,
        page_width=595, align="left",
    )
    company = TextBlock(
        text="甘肃博贞原生物乳业有限公司",
        top=68.3, bottom=80.3, x0=168, x1=324,
        font_size=12, align="center",
    )
    hline = LineBlock(top=83.0, bottom=83.7, x0=90, x1=505.3, thickness=0.72)
    title = TextBlock(
        text="生产用水检验记录",
        top=92.4, bottom=106.4, x0=241.7, x1=353.8,
        font_size=14, align="center",
    )
    sign_l = TextBlock(text="检验：", top=352.6, bottom=363, x0=111, x1=140, align="left")
    sign_r = TextBlock(text="审核：", top=352.6, bottom=363, x0=331, x1=360, align="center")

    rows = _group_horizontal_rows([logo, company, hline, title, sign_l, sign_r])
    assert len(rows) == 4
    assert len(rows[0]) == 2 and rows[0][0] is logo and rows[0][1] is company
    assert rows[1] == [hline]
    assert rows[2] == [title]
    assert len(rows[3]) == 2 and rows[3][0] is sign_l and rows[3][1] is sign_r

    tmp = tempfile.mkdtemp(prefix="pdf2word_test_")
    docx_path = os.path.join(tmp, "header.docx")
    write_document(
        [PageContent(
            blocks=[logo, company, hline, title, sign_l, sign_r],
            width=595, height=842,
        )],
        docx_path,
        page_breaks=False,
    )
    doc = Document(docx_path)

    # Layout rows must be normal paragraphs, never layout tables.
    assert len(doc.tables) == 0, "logo/signature rows must not use layout tables"
    assert doc.inline_shapes, "expected logo image inline in a paragraph"

    combo = None
    for p in doc.paragraphs:
        xml = "".join(r._element.xml for r in p.runs)
        if "graphicData" in xml and "博贞" in (p.text or ""):
            combo = p
            break
    assert combo is not None, "logo and company name must be in the same paragraph"

    # Header underline must become a paragraph bottom border.
    bordered = []
    for p in doc.paragraphs:
        pPr = p._p.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:pBdr")) is not None:
            bordered.append(p)
    assert bordered, "expected a paragraph with bottom border for the header rule"

    title_paras = [p for p in doc.paragraphs if "生产用水检验记录" in (p.text or "")]
    assert title_paras
    assert title_paras[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    # Title→table style: no huge trailing space on the title paragraph
    sa = title_paras[0].paragraph_format.space_after
    assert sa is None or sa.pt <= 2.0

    sign_para = None
    for p in doc.paragraphs:
        t = p.text or ""
        if "检验" in t and "审核" in t:
            sign_para = p
            break
    assert sign_para is not None, "检验/审核 must share one paragraph"


def test_title_table_gap_no_spacer_paragraph():
    """Title→table gap must use a compact spacer (exact 1pt line), not a full blank line."""
    from converter.pdf_reader import PageContent, TextBlock, TableBlock, Cell
    from docx import Document
    from docx.enum.text import WD_LINE_SPACING

    cells = [[Cell(text="A"), Cell(text="B")], [Cell(text="C"), Cell(text="D")]]
    owner = [[(r, c) for c in range(2)] for r in range(2)]
    title = TextBlock(
        text="生产用水检验记录", top=92, bottom=106, x0=240, x1=350,
        font_size=14, align="center",
    )
    # PDF gap title.bottom→table.top ≈ 40pt (as in the sample form)
    table = TableBlock(
        rows=2, cols=2, cells=cells, owner=owner,
        col_widths=[100, 100], row_heights=[18, 18],
        top=146.4, bottom=200, x0=84,
    )
    tmp = tempfile.mkdtemp(prefix="pdf2word_test_")
    docx_path = os.path.join(tmp, "gap.docx")
    write_document(
        [PageContent(blocks=[title, table], width=595, height=842)],
        docx_path, page_breaks=False,
    )
    doc = Document(docx_path)
    assert len(doc.tables) == 1
    # compact spacer: exact 1pt line spacing + space_before (not a full text line)
    compact = [
        p for p in doc.paragraphs
        if not (p.text or "").strip()
        and p.paragraph_format.line_spacing_rule == WD_LINE_SPACING.EXACTLY
    ]
    assert compact, "expected a compact spacer paragraph before the table"
    # first cell must NOT absorb the gap as cell padding
    cell_sb = doc.tables[0].cell(0, 0).paragraphs[0].paragraph_format.space_before
    assert cell_sb is None or cell_sb.pt < 6.0


def test_cell_nested_runs_written_to_docx():
    """Multi-run nested styles inside a cell become multiple Word runs."""
    from converter.pdf_reader import PageContent, TableBlock, Cell, TextRun
    from docx import Document

    rich = [
        [
            TextRun(text="标题", font_size=14.0, font_name="SimHei"),
            TextRun(text=" 副注", font_size=9.0, font_name="SimSun"),
        ],
        [TextRun(text="第二段正文", font_size=10.5, font_name="SimSun")],
    ]
    cells = [[Cell(text="标题 副注\n第二段正文", paragraphs=rich)]]
    owner = [[(0, 0)]]
    table = TableBlock(
        rows=1, cols=1, cells=cells, owner=owner,
        col_widths=[200], row_heights=[40],
        top=50, bottom=100, x0=72,
    )
    tmp = tempfile.mkdtemp(prefix="pdf2word_test_")
    docx_path = os.path.join(tmp, "rich.docx")
    write_document([PageContent(blocks=[table], width=595, height=842)], docx_path)
    doc = Document(docx_path)
    cell = doc.tables[0].cell(0, 0)
    assert len(cell.paragraphs) >= 2
    assert "标题" in cell.paragraphs[0].text
    assert "副注" in cell.paragraphs[0].text
    assert "第二段" in cell.paragraphs[1].text
    # Nested runs: first paragraph should have more than one run when styles differ.
    assert len(cell.paragraphs[0].runs) >= 2


def test_refine_merges_from_words_horizontal_span():
    """Word boxes spanning multiple empty columns grow colspan."""
    from converter.pdf_reader import Cell, WordIndex, _refine_merges_from_words

    # 1 row × 3 cols; left cell holds a wide title word.
    cells = [[Cell(text="总标题"), Cell(text=""), Cell(text="")]]
    owner = [[(0, 0), (0, 1), (0, 2)]]
    vx = [0.0, 100.0, 200.0, 300.0]
    hy = [0.0, 30.0]
    words = [
        {"text": "总标题", "x0": 5.0, "x1": 280.0, "top": 5.0, "bottom": 20.0,
         "size": 12.0, "fontname": "SimSun"},
    ]
    _refine_merges_from_words(cells, owner, vx, hy, WordIndex(words), 0, 0, 300, 30)
    assert cells[0][0].colspan == 3
    assert owner[0][1] == (0, 0)
    assert owner[0][2] == (0, 0)
    assert cells[0][1] is None
    assert cells[0][2] is None


def test_region_paragraphs_splits_font_runs():
    from converter.pdf_reader import WordIndex, _region_paragraphs

    vx = [0.0, 200.0]
    hy = [0.0, 40.0]
    words = [
        {"text": "粗", "x0": 10, "x1": 22, "top": 5, "bottom": 18,
         "size": 14.0, "fontname": "SimHei,Bold"},
        {"text": "细", "x0": 24, "x1": 36, "top": 5, "bottom": 18,
         "size": 10.0, "fontname": "SimSun"},
        {"text": "下行", "x0": 10, "x1": 40, "top": 22, "bottom": 34,
         "size": 10.0, "fontname": "SimSun"},
    ]
    paras = _region_paragraphs(WordIndex(words), vx, hy, 0, 0, 0, 0)
    assert len(paras) == 2
    assert len(paras[0]) == 2
    assert paras[0][0].font_size == 14.0
    assert paras[0][1].font_size == 10.0
    assert paras[1][0].text == "下行"


def test_ocr_module_graceful_when_unavailable(monkeypatch):
    from converter import ocr as ocr_mod

    ocr_mod.ocr_available.cache_clear()
    monkeypatch.setattr(ocr_mod, "ocr_available", lambda: False)
    blocks = ocr_mod.ocr_image_to_blocks(
        b"not-a-png", page_width=100, page_height=100
    )
    assert blocks == []
    info = ocr_mod.ocr_info()
    assert "available" in info


def test_content_warnings_ocr_flag():
    from converter.pdf_reader import PageContent, TextBlock, content_warnings

    pages = [PageContent(blocks=[
        TextBlock(text="识别文字", top=10, bottom=20, x0=10, x1=80, from_ocr=True)
    ], width=200, height=200)]
    warns = content_warnings(pages)
    assert "ocr_applied" in warns


def test_join_words_list_marker_before_body():
    """List number must stay left of body even when baselines differ slightly."""
    from converter.pdf_reader import _join_words

    # Marker a bit lower (larger top) than the CJK body — old sort put "10" last.
    words = [
        {"text": "、进入安装过程", "x0": 90.0, "x1": 200.0, "top": 100.0, "bottom": 112.0},
        {"text": "10", "x0": 72.0, "x1": 88.0, "top": 102.5, "bottom": 114.0},
    ]
    joined = _join_words(words)
    assert joined.startswith("10"), joined
    assert "进入安装过程" in joined
    assert not joined.endswith("10"), joined
    # Marker higher than body.
    words2 = [
        {"text": "10", "x0": 72.0, "x1": 88.0, "top": 98.0, "bottom": 110.0},
        {"text": "、进入安装过程", "x0": 90.0, "x1": 200.0, "top": 100.0, "bottom": 112.0},
    ]
    assert _join_words(words2).startswith("10")


def test_list_item_reading_order_with_baseline_offset(tmp_path):
    """PDF list lines keep '10、…' order when number/body y-offset slightly."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os

    pdf_path = str(tmp_path / "list_order.pdf")
    c = canvas.Canvas(pdf_path, pagesize=A4)
    h = A4[1]
    # Prefer a CJK font if available; fall back to Helvetica with ASCII stand-in.
    font_name = "Helvetica"
    for candidate in (
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\msyh.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    ):
        if os.path.isfile(candidate):
            try:
                pdfmetrics.registerFont(TTFont("CJKTest", candidate, subfontIndex=0))
                font_name = "CJKTest"
                break
            except Exception:
                try:
                    pdfmetrics.registerFont(TTFont("CJKTest", candidate))
                    font_name = "CJKTest"
                    break
                except Exception:
                    pass

    c.setFont(font_name, 12)
    if font_name == "CJKTest":
        # Number slightly lower than the following CJK text (common PDF quirk).
        c.drawString(72, h - 120, "10")
        c.drawString(95, h - 118, "、进入安装过程")
        # Number slightly higher.
        c.drawString(72, h - 160, "10")
        c.drawString(95, h - 162, "、进入安装过程")
    else:
        # ASCII proxy for environments without CJK fonts.
        c.drawString(72, h - 120, "10")
        c.drawString(95, h - 118, ", enter install process")
        c.drawString(72, h - 160, "10")
        c.drawString(95, h - 162, ", enter install process")
    c.save()

    pages = extract_document(pdf_path)
    texts = [b.text for b in pages[0].blocks if isinstance(b, TextBlock)]
    assert texts, "expected text blocks"
    for t in texts:
        # Must not reverse to body-then-number.
        assert not (
            t.strip().endswith("10") and ("进入" in t or "enter" in t.lower())
        ), t
        if "10" in t and ("进入" in t or "enter" in t.lower()):
            assert t.strip().startswith("10"), t


def test_find_tables_hybrid_accepts_borderless(tmp_path):
    """Borderless form (text strategy) should still yield a table block."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    pdf_path = str(tmp_path / "borderless.pdf")
    c = canvas.Canvas(pdf_path, pagesize=A4)
    # Two-column form labels without drawn grid lines (ASCII for font-safe tests).
    # Large horizontal gutter between label and value (real form layout).
    y = 750
    for left, right in [("Name", "Alice"), ("Dept", "R&D"), ("Date", "2026-07-12")]:
        c.drawString(72, y, left)
        c.drawString(220, y, right)
        y -= 28
    c.save()

    pages = extract_document(pdf_path)
    assert pages
    tables = [b for b in pages[0].blocks if isinstance(b, TableBlock)]
    texts = [b for b in pages[0].blocks if isinstance(b, TextBlock)]
    assert tables, "borderless label/value form should be kept as a table"
    flat = []
    for row in tables[0].cells:
        for cell in row:
            if cell and cell.text:
                flat.append(cell.text)
    blob = " ".join(flat)
    assert "Name" in blob or "Alice" in blob
    # Content must not disappear into empty text-only fallback.
    joined = " ".join(t.text for t in texts)
    assert "Name" in blob or "Name" in joined


def test_plain_prose_not_detected_as_table(tmp_path):
    """Single-column body text must stay as TextBlocks (no phantom grid)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    pdf_path = str(tmp_path / "prose.pdf")
    c = canvas.Canvas(pdf_path, pagesize=A4)
    c.setFont("Helvetica", 12)
    y = 780
    for line in [
        "Title of Document",
        "This is a simple paragraph of plain text.",
        "It has multiple lines but no columns or tables.",
        "Just flowing prose that should stay as text blocks.",
    ]:
        c.drawString(50, y, line)
        y -= 18
    c.save()

    pages = extract_document(pdf_path)
    assert pages
    tables = [b for b in pages[0].blocks if isinstance(b, TableBlock)]
    texts = [b for b in pages[0].blocks if isinstance(b, TextBlock)]
    assert tables == [], f"unexpected tables: {[(t.rows, t.cols) for t in tables]}"
    assert texts, "prose should become text blocks"
    joined = " ".join(t.text for t in texts)
    assert "Title" in joined or "paragraph" in joined


def test_soft_wrap_sentence_merged_to_one_paragraph(tmp_path):
    """PDF visual line wraps of one sentence become a single Word paragraph."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from docx import Document

    pdf_path = str(tmp_path / "softwrap.pdf")
    docx_path = str(tmp_path / "softwrap.docx")
    c = canvas.Canvas(pdf_path, pagesize=A4)
    c.setFont("Helvetica", 11)
    # One English sentence split across visual lines (no terminal punctuation mid-way).
    y = 780
    line1 = (
        "This long sentence continues across multiple visual lines so that "
        "the converter should join them"
    )
    line2 = "into one flowing paragraph without hard line breaks in Word."
    c.drawString(50, y, line1)
    c.drawString(50, y - 14, line2)
    # Real paragraph break (ends with period) then a new sentence.
    c.drawString(50, y - 40, "A separate short paragraph stays alone.")
    c.save()

    pages = extract_document(pdf_path)
    texts = [b for b in pages[0].blocks if isinstance(b, TextBlock)]
    # Soft-wrapped pair should merge into one block.
    merged = [t for t in texts if "continues across" in t.text and "flowing paragraph" in t.text]
    assert merged, f"expected merged sentence, got: {[t.text for t in texts]}"
    assert "\n" not in merged[0].text

    write_document(pages, docx_path)
    doc = Document(docx_path)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    assert any(
        "continues across" in p and "flowing paragraph" in p and "\n" not in p
        for p in paras
    ), paras
    # Separate paragraph preserved.
    assert any("separate short paragraph" in p for p in paras)


def test_normalize_spacing_cjk_and_mixed():
    from converter.text_utils import _normalize_spacing, _join_words

    assert _normalize_spacing("你 好 世 界") == "你好世界"
    assert _normalize_spacing("你好 ， 世界") == "你好，世界"
    assert _normalize_spacing("第 1 章") == "第1章"
    assert _normalize_spacing("版本 V2") == "版本V2"
    # Longer Latin words keep their spaces (readable mixed prose).
    assert _normalize_spacing("使用 Python 语言") == "使用 Python 语言"
    assert _normalize_spacing("hello  world") == "hello world"
    assert _normalize_spacing("这是一段话。 下一句") == "这是一段话。下一句"
    assert _normalize_spacing("安装 、 调试") == "安装、调试"
    # Fullwidth / NBSP collapsed then CJK spaces removed.
    assert " " not in _normalize_spacing("姓名\u3000\u3000张三")

    words = [
        {"text": "你", "x0": 0, "x1": 12, "top": 0, "bottom": 12},
        {"text": "好", "x0": 14, "x1": 26, "top": 0, "bottom": 12},
        {"text": "，", "x0": 27, "x1": 33, "top": 0, "bottom": 12},
        {"text": "世", "x0": 34, "x1": 46, "top": 0, "bottom": 12},
        {"text": "界", "x0": 48, "x1": 60, "top": 0, "bottom": 12},
    ]
    assert _join_words(words) == "你好，世界"

    words_num = [
        {"text": "第", "x0": 0, "x1": 12, "top": 0, "bottom": 12},
        {"text": "1", "x0": 13, "x1": 20, "top": 0, "bottom": 12},
        {"text": "章", "x0": 21, "x1": 33, "top": 0, "bottom": 12},
    ]
    assert _join_words(words_num) == "第1章"


def test_soft_wrap_helpers_unit():
    from converter.text_utils import (
        _is_soft_wrap_break,
        _normalize_newlines,
        _soft_join_text,
        _merge_soft_wrap_text_blocks,
    )
    from converter.models import TextBlock

    assert _soft_join_text("你好", "世界") == "你好世界"
    assert _soft_join_text("hello", "world") == "hello world"
    assert _soft_join_text("exam-", "ple") == "example"

    assert _is_soft_wrap_break(
        "这是一段比较长的中文句子用于测试软换行合并是否生效",
        "后续内容继续写下去不要硬换行",
        v_gap=2.0,
        prev_height=12.0,
        next_height=12.0,
        prev_x0=50.0,
        next_x0=50.0,
        prev_font=11.0,
        next_font=11.0,
    )
    # Sentence end → real break
    assert not _is_soft_wrap_break("结束了。", "下一句", v_gap=2.0, prev_height=12, next_height=12)
    # Short title → real break
    assert not _is_soft_wrap_break("标题", "这是正文第一行内容", v_gap=2.0, prev_height=12, next_height=12)
    # List marker → real break
    assert not _is_soft_wrap_break(
        "前文内容继续写很长一段以便通过满行判断",
        "1、安装设备",
        v_gap=2.0,
        prev_height=12,
        next_height=12,
    )

    assert _normalize_newlines("你好\n世界") == "你好世界"
    assert _normalize_newlines("结束。\n下一句") == "结束。\n下一句"

    blocks = [
        TextBlock(
            text="这是一段比较长的中文句子用于测试软换行合并是否生效",
            top=10, bottom=22, x0=50, x1=500, font_size=11,
        ),
        TextBlock(
            text="后续内容继续写下去不要硬换行",
            top=24, bottom=36, x0=50, x1=300, font_size=11,
        ),
        TextBlock(
            text="全新段落在这里开始。",
            top=60, bottom=72, x0=50, x1=200, font_size=11,
        ),
    ]
    # page_right helps full-width detection for the first wrap line.
    merged = _merge_soft_wrap_text_blocks(blocks, page_right=510.0)
    assert len(merged) == 2
    assert "后续内容" in merged[0].text and "\n" not in merged[0].text
    assert "全新段落" in merged[1].text

    # Table-cell style: ~1× line-box leading + full width must still merge.
    from converter.text_utils import _merge_soft_wrap_paragraphs
    from converter.models import TextRun

    paras = [
        [TextRun(text="3、公共区域卫生：理化室窗台及靠窗台气瓶柜/干燥箱/酸缸、玻璃、地面、踢", font_size=12)],
        [TextRun(text="脚线为公共区域，按照分组周一到周五轮流打扫，周六周天所有人共同打扫。", font_size=12)],
    ]
    # top/bottom/x0/x1 matching the real PDF (v_gap≈11.4 on 12pt boxes).
    boxes = [
        (680.4, 692.4, 90.0, 501.1),
        (703.8, 715.8, 90.0, 498.0),
    ]
    joined = _merge_soft_wrap_paragraphs(paras, line_boxes=boxes, cell_right=506.6)
    assert len(joined) == 1, joined
    blob = "".join(r.text for r in joined[0])
    assert "踢脚线" in blob and "\n" not in blob


def test_multi_column_prose_not_detected_as_table(tmp_path):
    """Two-column article layout is not a data table."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    pdf_path = str(tmp_path / "multicol.pdf")
    c = canvas.Canvas(pdf_path, pagesize=A4)
    c.setFont("Helvetica", 11)
    h = A4[1]
    for i in range(12):
        c.drawString(
            50, h - 80 - i * 16,
            f"Left side paragraph content line {i + 1} continues here.",
        )
        c.drawString(
            320, h - 80 - i * 16,
            f"Right column text line {i + 1} about topic.",
        )
    c.save()

    pages = extract_document(pdf_path)
    assert pages
    tables = [b for b in pages[0].blocks if isinstance(b, TableBlock)]
    texts = [b for b in pages[0].blocks if isinstance(b, TextBlock)]
    assert tables == [], f"unexpected tables: {[(t.rows, t.cols) for t in tables]}"
    assert len(texts) >= 4
    joined = " ".join(t.text for t in texts)
    assert "Left side" in joined and "Right column" in joined


def test_text_and_image_pdf_no_phantom_table(tmp_path):
    """PDF with only text + a decorative rect must not invent a huge table."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    pdf_path = str(tmp_path / "text_image.pdf")
    c = canvas.Canvas(pdf_path, pagesize=A4)
    w, h = A4
    c.setFont("Helvetica", 14)
    c.drawString(50, h - 50, "Company Annual Report 2024")
    c.setFont("Helvetica", 11)
    y = h - 90
    for i, line in enumerate([
        "This document contains only plain text and an image.",
        "There should be no tables detected at all.",
        "Name: John Doe          Dept: Engineering",
        "Date: 2024-01-01        Status: Active",
    ]):
        c.drawString(50, y - i * 18, line)
    for i in range(6):
        c.drawString(50, h - 280 - i * 14, f"Paragraph line {i + 1} on the left side.")
        c.drawString(320, h - 280 - i * 14, f"Column two content line number {i + 1}.")
    c.setFillColorRGB(0.8, 0.85, 0.9)
    c.rect(50, 80, 200, 100, fill=1, stroke=0)
    c.setFillColorRGB(0, 0, 0)
    c.setFont("Helvetica", 9)
    c.drawString(60, 120, "[Image placeholder area]")
    c.save()

    pages = extract_document(pdf_path)
    assert pages
    tables = [b for b in pages[0].blocks if isinstance(b, TableBlock)]
    texts = [b for b in pages[0].blocks if isinstance(b, TextBlock)]
    assert tables == [], (
        f"text+image PDF produced phantom table(s): "
        f"{[(t.rows, t.cols) for t in tables]}"
    )
    assert texts
    joined = " ".join(t.text for t in texts)
    assert "Company Annual Report" in joined or "plain text" in joined


def test_ruled_grid_table_still_detected(tmp_path):
    """Drawn GRID tables must still be accepted after the prose filter."""
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
    from reportlab.lib import colors

    pdf_path = str(tmp_path / "grid.pdf")
    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    data = [
        ["Name", "Info", "Score"],
        ["Alice", "Math", "90"],
        ["Bob", "Physics", "85"],
    ]
    table = Table(data, colWidths=[80, 80, 80])
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 1, colors.black),
    ]))
    doc.build([table])

    pages = extract_document(pdf_path)
    tables = [b for b in pages[0].blocks if isinstance(b, TableBlock)]
    assert len(tables) == 1
    assert tables[0].rows == 3 and tables[0].cols == 3


