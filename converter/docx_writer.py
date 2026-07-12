from __future__ import annotations

import io
from typing import List, Optional, Sequence, Tuple, Union

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from .pdf_reader import (
    PageContent,
    TextBlock,
    TableBlock,
    ImageBlock,
    LineBlock,
    Cell,
    TextRun,
)

# ----- style defaults -------------------------------------------------------
DEFAULT_FONT = "宋体"
DEFAULT_SIZE = 10.5
# Fallback page margins when content bounds cannot be inferred (inches).
DEFAULT_MARGIN_IN = 0.7
# Vertical tolerance (pt) for grouping blocks that sit on the same visual row.
ROW_Y_TOL = 10.0
# Compact spacer height before a table that follows a title (pt).
TABLE_SPACER_PT = 1.0
# Soft vertical gaps between other rows: ignore tiny PDF gaps, cap large ones.
V_GAP_MIN_PT = 6.0
V_GAP_CAP_PT = 28.0

_HALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}
_VALIGN = {
    "top": WD_ALIGN_VERTICAL.TOP,
    "center": WD_ALIGN_VERTICAL.CENTER,
    "bottom": WD_ALIGN_VERTICAL.BOTTOM,
}

# Map common embedded PDF font names to the Word font that renders them.
_FONT_MAP = {
    "simsun": "宋体",
    "simhei": "黑体",
    "microsoftyahei": "微软雅黑",
    "stsong": "华文宋体",
    "stxihei": "华文细黑",
    "stkaiti": "华文楷体",
    "kaiti": "楷体",
    "fangsong": "仿宋",
}

Block = Union[TextBlock, TableBlock, ImageBlock, LineBlock]


def _normalize_font(name: Optional[str]) -> str:
    if not name:
        return DEFAULT_FONT
    # strip embedded-subset prefix, e.g. "ABCDEF+SimSun" -> "SimSun"
    if "+" in name:
        name = name.split("+", 1)[1]
    lower = name.lower()
    for key, value in _FONT_MAP.items():
        if key in lower:
            return value
    return name


def _is_bold(font_name: Optional[str]) -> bool:
    return bool(font_name) and "bold" in font_name.lower()


def _set_run_font(run, font_name: Optional[str], font_size: Optional[float],
                  bold: bool) -> None:
    name = _normalize_font(font_name)
    run.font.name = name
    run.font.size = Pt(font_size) if font_size else Pt(DEFAULT_SIZE)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def _set_cell_text(cell, text: str, font_name: Optional[str],
                   font_size: Optional[float], bold: bool) -> None:
    lines = (text or "").split("\n")
    # first line via cell.text (creates the primary run); remaining lines are
    # appended after explicit line breaks so intra-cell newlines are preserved.
    cell.text = lines[0] if lines else ""
    paragraph = cell.paragraphs[0]
    for run in paragraph.runs:
        _set_run_font(run, font_name, font_size, bold)
    for line in lines[1:]:
        br = paragraph.add_run("")
        br.add_break()
        run = paragraph.add_run(line)
        _set_run_font(run, font_name, font_size, bold)


def _set_cell_rich(
    cell,
    paragraphs: List[List[TextRun]],
    *,
    fallback_font: Optional[str] = None,
    fallback_size: Optional[float] = None,
    align: str = "left",
) -> None:
    """Write multi-paragraph / multi-run nested styles into a table cell.

    Each inner list is one paragraph; each ``TextRun`` becomes a Word run with
    its own font name/size (bold inferred from the PDF font name).
    """
    if not paragraphs:
        cell.text = ""
        return

    # Clear default empty paragraph content.
    cell.text = ""
    # Reuse the first paragraph; add more as needed.
    first = True
    for para_runs in paragraphs:
        if first:
            p = cell.paragraphs[0]
            # Remove any default run left by cell.text = "".
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            first = False
        else:
            p = cell.add_paragraph()
        p.alignment = _HALIGN.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if not para_runs:
            continue
        for tr in para_runs:
            if not tr.text:
                continue
            run = p.add_run(tr.text)
            fname = tr.font_name or fallback_font
            fsize = tr.font_size if tr.font_size is not None else fallback_size
            _set_run_font(run, fname, fsize, _is_bold(fname))


def _set_cell_borders(cell, borders: Optional[dict]) -> None:
    if not borders:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_b = tc_pr.find(qn("w:tcBorders"))
    if tc_b is None:
        tc_b = OxmlElement("w:tcBorders")
        tc_pr.append(tc_b)
    else:
        tc_b.clear()
    for edge in ("top", "left", "bottom", "right"):
        val = borders.get(edge)
        if not val:
            continue
        w_pt, color, dashed = val
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "dashed" if dashed else "single")
        el.set(qn("w:sz"), str(int(round(w_pt * 8))))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tc_b.append(el)


def _set_table_borders(table, outer_pt: float, inner_pt: float,
                       color: str, dashed: bool) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    else:
        borders.clear()
    val = "dashed" if dashed else "single"
    edges = {
        "top": outer_pt, "left": outer_pt, "bottom": outer_pt, "right": outer_pt,
        "insideH": inner_pt, "insideV": inner_pt,
    }
    for name, pt in edges.items():
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), str(int(round(pt * 8))))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)


def _set_cell_shading(cell, hex_color: str) -> None:
    if not hex_color:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


def _set_fixed_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _set_table_dims(table, col_widths: List[float], row_heights: List[float]) -> None:
    if col_widths:
        table.width = Pt(sum(col_widths))
        for i, col in enumerate(table.columns):
            if i < len(col_widths):
                col.width = Pt(col_widths[i])
    if row_heights:
        for i, row in enumerate(table.rows):
            if i < len(row_heights):
                row.height = Pt(row_heights[i])
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    _set_fixed_layout(table)


def _set_table_left_indent(table, indent_pt: float) -> None:
    """Shift a table from the left page margin by ``indent_pt`` points."""
    if indent_pt <= 0:
        return
    tbl_pr = table._tbl.tblPr
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    # twips = pt * 20
    ind.set(qn("w:w"), str(int(round(indent_pt * 20))))
    ind.set(qn("w:type"), "dxa")


def _block_x0(block: Block) -> float:
    if isinstance(block, ImageBlock):
        return float(block.x0 or 0.0)
    if isinstance(block, LineBlock):
        return float(block.x0 or 0.0)
    if isinstance(block, TableBlock):
        return float(block.x0 or 0.0)
    return float(getattr(block, "x0", 0.0) or 0.0)


def _block_x1(block: Block) -> float:
    if isinstance(block, ImageBlock):
        return float(block.x0 or 0.0) + float(block.width_pt or 0.0)
    if isinstance(block, LineBlock):
        return float(block.x1 or 0.0)
    if isinstance(block, TableBlock):
        w = sum(block.col_widths) if block.col_widths else 0.0
        return float(block.x0 or 0.0) + w
    return float(getattr(block, "x1", 0.0) or 0.0)


def _block_top(block: Block) -> float:
    return float(getattr(block, "top", 0.0) or 0.0)


def _block_bottom(block: Block) -> float:
    return float(getattr(block, "bottom", 0.0) or 0.0)


def _infer_page_margins(
    pages: Sequence[PageContent],
) -> Tuple[float, float, float, float, float, float]:
    """Return (page_w, page_h, left, right, top, bottom) in points.

    Margins are inferred from the union content bounding box when possible so
    that PDF x-coordinates map 1:1 into the Word content area. Falls back to
    DEFAULT_MARGIN_IN when bounds are missing.
    """
    page_w = 0.0
    page_h = 0.0
    for page in pages:
        page_w = max(page_w, float(page.width or 0.0))
        page_h = max(page_h, float(page.height or 0.0))

    # A4 defaults if page size unknown
    if page_w <= 0:
        page_w = 595.27
    if page_h <= 0:
        page_h = 841.89

    xs0: List[float] = []
    xs1: List[float] = []
    ys0: List[float] = []
    ys1: List[float] = []
    for page in pages:
        for b in page.blocks:
            if isinstance(b, TableBlock) and not b.cells:
                continue
            x0, x1 = _block_x0(b), _block_x1(b)
            y0, y1 = _block_top(b), _block_bottom(b)
            if x1 > x0:
                xs0.append(x0)
                xs1.append(x1)
            if y1 > y0:
                ys0.append(y0)
                ys1.append(y1)

    default_m = DEFAULT_MARGIN_IN * 72.0  # pt
    if not xs0:
        return page_w, page_h, default_m, default_m, default_m, default_m

    content_left = min(xs0)
    content_right = max(xs1)
    content_top = min(ys0) if ys0 else default_m
    content_bottom = max(ys1) if ys1 else page_h - default_m

    # Use the content inset as the Word margin so left-aligned content at the
    # content edge needs no extra indent. Clamp to a reasonable range.
    left = max(36.0, min(content_left, 144.0))
    right = max(36.0, min(page_w - content_right, 144.0))
    # Keep left/right roughly symmetric when the PDF is balanced (forms often are).
    if abs(left - right) < 12.0:
        left = right = (left + right) / 2.0
    top = max(36.0, min(content_top, 144.0))
    bottom = max(36.0, min(page_h - content_bottom, 144.0))

    # When page dimensions were not provided on the PageContent (tests often
    # omit them), prefer the classic 0.7" default so indent math stays stable.
    any_size = any(float(p.width or 0) > 0 and float(p.height or 0) > 0 for p in pages)
    if not any_size:
        m = default_m
        return page_w, page_h, m, m, m, m

    return page_w, page_h, left, right, top, bottom


def _pdf_x_to_indent_pt(pdf_x: float, left_margin_pt: float) -> float:
    """Extra left indent (pt) so content at ``pdf_x`` lands at the right place."""
    return max(0.0, float(pdf_x) - float(left_margin_pt))


def _apply_section_geometry(
    doc: Document,
    page_w: float,
    page_h: float,
    left: float,
    right: float,
    top: float,
    bottom: float,
) -> None:
    section = doc.sections[0]
    section.page_width = Pt(page_w)
    section.page_height = Pt(page_h)
    section.left_margin = Pt(left)
    section.right_margin = Pt(right)
    section.top_margin = Pt(top)
    section.bottom_margin = Pt(bottom)


def _group_horizontal_rows(blocks: Sequence[Block]) -> List[List[Block]]:
    """Group blocks that share a vertical band into visual rows.

    Tables always form their own row. Other blocks (text / image / line) are
    merged when their vertical centres (or overlapping bands) fall within
    ``ROW_Y_TOL``. Within a row, blocks are ordered left-to-right by ``x0``.
    """
    if not blocks:
        return []

    # Keep source order for primary sort, but group by vertical position.
    items = list(blocks)
    rows: List[List[Block]] = []
    used = [False] * len(items)

    for i, b in enumerate(items):
        if used[i]:
            continue
        if isinstance(b, TableBlock):
            used[i] = True
            rows.append([b])
            continue

        # Seed a non-table row with this block.
        band: List[Block] = [b]
        used[i] = True
        top = _block_top(b)
        bottom = _block_bottom(b) or top
        cy = (top + bottom) / 2.0

        # Absorb later non-table blocks whose centre is close to this band.
        changed = True
        while changed:
            changed = False
            for j, other in enumerate(items):
                if used[j] or isinstance(other, TableBlock):
                    continue
                ot = _block_top(other)
                ob = _block_bottom(other) or ot
                ocy = (ot + ob) / 2.0
                # Same row if centres are close OR the vertical ranges overlap
                # substantially (logo + company name often slightly offset).
                close = abs(ocy - cy) <= ROW_Y_TOL
                overlap = not (ob < top - 1 or ot > bottom + 1)
                if close or (overlap and abs(ocy - cy) <= ROW_Y_TOL * 2):
                    band.append(other)
                    used[j] = True
                    top = min(top, ot)
                    bottom = max(bottom, ob)
                    cy = (top + bottom) / 2.0
                    changed = True

        band.sort(key=_block_x0)
        rows.append(band)

    return rows


def _set_paragraph_spacing(paragraph, before: float = 0.0, after: float = 0.0) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.0


def _add_picture_run(paragraph, image: ImageBlock) -> None:
    """Embed ``image`` as an inline picture in ``paragraph``."""
    stream = io.BytesIO(image.image_bytes)
    width = Pt(image.width_pt) if image.width_pt else None
    height = Pt(image.height_pt) if image.height_pt else None
    run = paragraph.add_run()
    try:
        if width and height:
            run.add_picture(stream, width=width, height=height)
        elif width:
            run.add_picture(stream, width=width)
        else:
            run.add_picture(stream)
    except Exception:
        # Corrupt / non-PNG bytes — skip the picture rather than aborting.
        pass


def _write_text_paragraph(
    doc: Document,
    block: TextBlock,
    left_margin_pt: float,
    *,
    space_before_pt: float = 0.0,
) -> None:
    p = doc.add_paragraph()
    p.alignment = _HALIGN.get(block.align, WD_ALIGN_PARAGRAPH.LEFT)
    _set_paragraph_spacing(p, before=space_before_pt, after=0.0)
    if block.align == "left":
        indent = _pdf_x_to_indent_pt(block.x0, left_margin_pt)
        if indent > 0.5:
            p.paragraph_format.left_indent = Pt(indent)
    run = p.add_run(block.text)
    _set_run_font(run, block.font_name, block.font_size, _is_bold(block.font_name))


def _write_image_paragraph(
    doc: Document,
    image: ImageBlock,
    left_margin_pt: float,
    *,
    space_before_pt: float = 0.0,
) -> None:
    p = doc.add_paragraph()
    align = image.align or "left"
    p.alignment = _HALIGN.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    _set_paragraph_spacing(p, before=space_before_pt, after=0.0)
    if align == "left":
        indent = _pdf_x_to_indent_pt(image.x0, left_margin_pt)
        if indent > 0.5:
            p.paragraph_format.left_indent = Pt(indent)
    _add_picture_run(p, image)


def _write_hline(
    doc: Document,
    line: LineBlock,
    left_margin_pt: float,
    page_width: float,
    *,
    space_before_pt: float = 0.0,
) -> None:
    """Render a horizontal rule as a paragraph bottom border (pBdr)."""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=space_before_pt, after=0.0)
    # Keep the rule height tight so it doesn't introduce a large gap.
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(max(line.thickness, 1.0))

    # Indent so the border starts near the rule's PDF x0.
    indent = _pdf_x_to_indent_pt(line.x0, left_margin_pt)
    if indent > 0.5:
        p.paragraph_format.left_indent = Pt(indent)

    # Optionally constrain width via right indent so the border doesn't span
    # the full content width (Word pBdr is full-paragraph, but right indent
    # shortens the paragraph box).
    content_width = max(page_width - 2 * left_margin_pt, 1.0)
    rule_width = max(float(line.x1) - float(line.x0), 1.0)
    # Approximate right indent relative to content area.
    right_extra = max(0.0, content_width - indent - rule_width)
    if right_extra > 1.0:
        p.paragraph_format.right_indent = Pt(right_extra)

    # Build <w:pBdr><w:bottom .../></w:pBdr>
    pPr = p._p.get_or_add_pPr()
    # Remove any existing pBdr
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    # Word border size is in eighths of a point.
    sz = max(2, int(round(max(line.thickness, 0.5) * 8)))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), line.color or "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _write_multi_block_row(
    doc: Document,
    row: Sequence[Block],
    left_margin_pt: float,
    page_width: float,
    *,
    space_before_pt: float = 0.0,
) -> None:
    """Write several side-by-side blocks as one paragraph with tab stops.

    Used for logo+company header and 检验/审核 signature labels. Never uses a
    layout table (tests assert ``len(doc.tables) == 0`` for those cases).
    """
    # Pure line row → hline only.
    if len(row) == 1 and isinstance(row[0], LineBlock):
        _write_hline(
            doc, row[0], left_margin_pt, page_width,
            space_before_pt=space_before_pt,
        )
        return
    if len(row) == 1 and isinstance(row[0], TextBlock):
        _write_text_paragraph(
            doc, row[0], left_margin_pt, space_before_pt=space_before_pt,
        )
        return
    if len(row) == 1 and isinstance(row[0], ImageBlock):
        _write_image_paragraph(
            doc, row[0], left_margin_pt, space_before_pt=space_before_pt,
        )
        return

    # Filter out LineBlocks that somehow share a band (rare) — write them alone.
    non_lines = [b for b in row if not isinstance(b, LineBlock)]
    lines = [b for b in row if isinstance(b, LineBlock)]
    if not non_lines:
        for i, lb in enumerate(lines):
            _write_hline(
                doc, lb, left_margin_pt, page_width,
                space_before_pt=space_before_pt if i == 0 else 0.0,
            )
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(p, before=space_before_pt, after=0.0)

    first_x0 = _block_x0(non_lines[0])
    base_indent = _pdf_x_to_indent_pt(first_x0, left_margin_pt)
    if base_indent > 0.5:
        p.paragraph_format.left_indent = Pt(base_indent)

    # Content-area width for tab-stop positions (relative to left margin).
    # Tab positions in python-docx are measured from the left page margin.
    tab_stops = p.paragraph_format.tab_stops
    # Clear any default stops by replacing with our own.
    for i, block in enumerate(non_lines):
        if i == 0:
            continue
        # Position of this block relative to the left margin (absolute from margin).
        pos_pt = _pdf_x_to_indent_pt(_block_x0(block), left_margin_pt)
        # Tab stop position is from the left margin, not from paragraph indent.
        # When left_indent is set, tab positions are still from the margin.
        if pos_pt > base_indent + 1.0:
            tab_stops.add_tab_stop(Pt(pos_pt), alignment=WD_TAB_ALIGNMENT.LEFT)

    for i, block in enumerate(non_lines):
        if i > 0:
            p.add_run("\t")
        if isinstance(block, ImageBlock):
            _add_picture_run(p, block)
        elif isinstance(block, TextBlock):
            run = p.add_run(block.text)
            _set_run_font(
                run, block.font_name, block.font_size, _is_bold(block.font_name)
            )

    for lb in lines:
        _write_hline(doc, lb, left_margin_pt, page_width)


def _add_compact_spacer(doc: Document, height_pt: float = TABLE_SPACER_PT) -> None:
    """Insert an exact-height empty paragraph (title→table / soft v-gap)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(max(height_pt, 0.5))


def _soft_v_gap_pt(gap_pt: float) -> float:
    """Map a PDF vertical gap to a modest Word spacer height."""
    if gap_pt < V_GAP_MIN_PT:
        return 0.0
    # Keep more of small gaps; compress large white space so forms stay compact.
    if gap_pt <= 18.0:
        return min(gap_pt, V_GAP_CAP_PT)
    return min(12.0 + (gap_pt - 18.0) * 0.25, V_GAP_CAP_PT)


def _row_top(row: Sequence[Block]) -> float:
    return min(_block_top(b) for b in row)


def _row_bottom(row: Sequence[Block]) -> float:
    return max((_block_bottom(b) or _block_top(b)) for b in row)


def _write_table(
    doc: Document,
    table: TableBlock,
    left_margin_pt: float,
    *,
    after_title: bool = False,
    pre_gap_pt: float = 0.0,
) -> None:
    if after_title:
        _add_compact_spacer(doc, TABLE_SPACER_PT)
    else:
        soft = _soft_v_gap_pt(pre_gap_pt)
        if soft > 0.5:
            _add_compact_spacer(doc, soft)

    grid = doc.add_table(rows=table.rows, cols=table.cols)
    try:
        grid.style = "Table Grid"
    except KeyError:
        pass

    has_cell_borders = False
    for r in range(table.rows):
        for c in range(table.cols):
            anchor = table.owner[r][c]
            if anchor != (r, c):
                continue  # covered cell of a merge; absorbed by the anchor
            cell_obj: Cell = table.cells[r][c]
            target = grid.cell(r, c)
            if cell_obj.rowspan > 1 or cell_obj.colspan > 1:
                end_r = min(r + cell_obj.rowspan - 1, table.rows - 1)
                end_c = min(c + cell_obj.colspan - 1, table.cols - 1)
                target = target.merge(grid.cell(end_r, end_c))
            if cell_obj.paragraphs:
                _set_cell_rich(
                    target,
                    cell_obj.paragraphs,
                    fallback_font=cell_obj.font_name,
                    fallback_size=cell_obj.font_size,
                    align=cell_obj.align,
                )
            else:
                _set_cell_text(
                    target, cell_obj.text,
                    font_name=cell_obj.font_name, font_size=cell_obj.font_size,
                    bold=_is_bold(cell_obj.font_name),
                )
            target.vertical_alignment = _VALIGN.get(
                cell_obj.valign, WD_ALIGN_VERTICAL.TOP
            )
            for paragraph in target.paragraphs:
                paragraph.alignment = _HALIGN.get(
                    cell_obj.align, WD_ALIGN_PARAGRAPH.LEFT
                )
                # Keep cell padding tight — do not absorb title→table gap here.
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
            _set_cell_shading(target, cell_obj.bg_color or "")
            if cell_obj.borders:
                has_cell_borders = True
                _set_cell_borders(target, cell_obj.borders)

    # Per-cell borders win when present (they already cover every drawn edge,
    # including merges). Otherwise fall back to a uniform grid so borderless /
    # text-strategy tables still render as a table.
    if not has_cell_borders:
        _set_table_borders(
            grid, table.border_outer, table.border_inner,
            table.border_color, table.border_dashed,
        )
    _set_table_dims(grid, table.col_widths, table.row_heights)

    indent = _pdf_x_to_indent_pt(table.x0, left_margin_pt)
    if indent > 0.5:
        _set_table_left_indent(grid, indent)


def _is_title_like(block: Block) -> bool:
    if not isinstance(block, TextBlock):
        return False
    if block.align == "center" and block.font_size and block.font_size >= 12:
        return True
    # Short centered-ish heading without needing font size
    return block.align == "center" and 4 <= len((block.text or "").strip()) <= 40


def write_document(
    pages: List[PageContent],
    output_path: str,
    *,
    page_breaks: bool = True,
) -> None:
    """Write extracted pages to a .docx file.

    When ``page_breaks`` is True (default), a page break is inserted between
    consecutive PDF pages so the Word document mirrors the source pagination.

    Layout notes
    ------------
    * Page size and margins are inferred from PDF page dimensions / content
      bounding boxes so horizontal positions match the source (avoids the
      classic "everything shifted left" problem from undersized margins).
    * Same-row items (logo + company, 检验/审核) share one paragraph with tab
      stops — never a layout table.
    * Horizontal rules become paragraph bottom borders (``pBdr``).
    * Title→table gaps use a compact 1pt spacer paragraph.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = DEFAULT_FONT
    style.font.size = Pt(DEFAULT_SIZE)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)

    page_w, page_h, left_m, right_m, top_m, bottom_m = _infer_page_margins(pages)
    _apply_section_geometry(doc, page_w, page_h, left_m, right_m, top_m, bottom_m)

    for i, page in enumerate(pages):
        if page_breaks and i > 0:
            doc.add_page_break()

        # Prefer per-page width when available.
        pw = float(page.width or page_w)
        rows = _group_horizontal_rows(list(page.blocks))
        prev_was_title = False
        prev_bottom: Optional[float] = None

        for row in rows:
            top = _row_top(row)
            gap = 0.0 if prev_bottom is None else max(0.0, top - prev_bottom)
            soft_before = _soft_v_gap_pt(gap) if prev_bottom is not None else 0.0

            if len(row) == 1 and isinstance(row[0], TableBlock):
                _write_table(
                    doc, row[0], left_m,
                    after_title=prev_was_title,
                    pre_gap_pt=0.0 if prev_was_title else gap,
                )
                prev_was_title = False
                prev_bottom = _row_bottom(row)
                continue

            if len(row) == 1 and isinstance(row[0], LineBlock):
                _write_hline(
                    doc, row[0], left_m, pw, space_before_pt=soft_before,
                )
                prev_was_title = False
                prev_bottom = _row_bottom(row)
                continue

            if len(row) == 1 and isinstance(row[0], ImageBlock):
                _write_image_paragraph(
                    doc, row[0], left_m, space_before_pt=soft_before,
                )
                prev_was_title = False
                prev_bottom = _row_bottom(row)
                continue

            if len(row) == 1 and isinstance(row[0], TextBlock):
                _write_text_paragraph(
                    doc, row[0], left_m, space_before_pt=soft_before,
                )
                # Title paragraphs should not leave large trailing space.
                last = doc.paragraphs[-1]
                last.paragraph_format.space_after = Pt(0)
                prev_was_title = _is_title_like(row[0])
                prev_bottom = _row_bottom(row)
                continue

            # Multi-block horizontal row (logo+text, dual signatures, …)
            _write_multi_block_row(
                doc, row, left_m, pw, space_before_pt=soft_before,
            )
            prev_was_title = False
            prev_bottom = _row_bottom(row)

    doc.save(output_path)
